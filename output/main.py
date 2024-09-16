import json
import os
import time
from copy import copy
from pathlib import Path
from types import GeneratorType
import subprocess
from io import BytesIO
import re
import streamlit as st
from PIL import Image
from utils import dict_iter_render_message, parse_image_content, parse_list_from_html, render_single_dict_message
from docx import Document

from promptflow._constants import STREAMING_ANIMATION_TIME
from promptflow._sdk._orchestrator import TestSubmitter
from promptflow._sdk._orchestrator.utils import resolve_generator, resolve_generator_output_with_cache
from promptflow._utils.flow_utils import dump_flow_result
from promptflow._utils.multimedia_utils import BasicMultimediaProcessor
from promptflow.client import load_flow


def run_command():
    command = [
        "pf", "connection", "create",
        "--file", "/workspaces/HV_Enigma/output/flow/azure_openai.yaml",
        "--set", "api_key=498b28b13a51434ab7c614bd4d961f20",
        "api_base=https://enigmaengine.openai.azure.com/",
        "--name", "azure_open_ai"
    ]
    try:
        subprocess.run(command, check=True)
    except FileNotFoundError:
        st.error("The command 'pf' was not found. Please ensure it is installed and in your PATH.")
    except subprocess.CalledProcessError as e:
        st.error(f"An error occurred while running the command: {e}")
    except Exception as e:
        st.error(f"An error occurred: {e}")


# def create_docx(content):
#     doc = Document()
#     doc.add_paragraph(content)
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer


def format_markdown(content):
    # Replace newlines with double newlines for markdown formatting
    content = content.replace("\\n", "\n\n")
    # Replace escaped characters
    content = content.replace("\\", "")
    return content


def start():
    def clear_chat() -> None:
        st.session_state.messages = []

    def render_message(role, message_items):
        if (role == 'user'):
            with st.chat_message(role):
                if is_chat_flow:
                    render_single_dict_message(message_items)
                else:
                    dict_iter_render_message(message_items)

    def show_conversation() -> None:
        if "messages" not in st.session_state:
            st.session_state.messages = []
            st.session_state.history = []
        if st.session_state.messages:
            for role, message_items in st.session_state.messages:
                render_message(role, message_items)

    def get_chat_history_from_session():
        if "history" in st.session_state:
            return st.session_state.history
        return []

    def post_process_dump_result(response, session_state_history, *, generator_record):
        response = resolve_generator(response, generator_record)
        # Get base64 for multi modal object
        # Just use BasicMultimediaProcessor to keep the original logic here.
        # TODO: Add support for other multimedia types
        multimedia_processor = BasicMultimediaProcessor()
        multimedia_response = {
            k: multimedia_processor.load_multimedia_data_recursively(v) for k, v in response.output.items()
        }
        resolved_outputs = {
            k: multimedia_processor.convert_multimedia_data_to_base64_dict(v) for k, v in multimedia_response.items()
        }
        st.session_state.messages.append(("assistant", resolved_outputs))
        session_state_history.update({"outputs": response.output})
        st.session_state.history.append(session_state_history)
        if is_chat_flow:
            dump_path = Path(flow_path).parent
            dump_flow_result(flow_folder=dump_path, flow_result=response, prefix="chat")
        
        # Save the output to a .txt file
        with open("/workspaces/HV_Enigma/output/hvenigma.txt", "w") as f:
            f.write(json.dumps(resolved_outputs, indent=4))
        
        return resolved_outputs

    def submit(**kwargs) -> None:
        # generator record should be reset for each submit
        generator_record = {}

        st.session_state.messages.append(("user", kwargs))
        session_state_history = dict()
        session_state_history.update({"inputs": kwargs})
        with container:
            render_message("user", kwargs)

        flow = load_flow(flow_path)
        with TestSubmitter(flow=flow, flow_context=flow.context).init(stream_output=is_streaming) as submitter:
            # can't exit the context manager before the generator is fully consumed
            response = submitter.flow_test(
                inputs={chat_history_input_name: get_chat_history_from_session(), **kwargs},
                allow_generator_output=is_streaming,
            )

            if response.run_info.status.value == "Failed":
                raise Exception(response.run_info.error)

            if is_streaming:
                # Display assistant response in chat message container
                with container:
                    with st.chat_message("assistant"):
                        message_placeholder = st.empty()
                        full_response = f"{chat_output_name}: "
                        full_response = re.sub(r'\\.',lambda x:{'\\n':'\n','\\t':'\t'}.get(x[0],x[0]),full_response)
                        full_response = format_markdown(full_response)
                        print(full_response)
                        prefix_length = len(full_response)
                        chat_output = response.output[chat_output_name]
                        if isinstance(chat_output, GeneratorType):
                            # Simulate stream of response with milliseconds delay
                            for chunk in resolve_generator_output_with_cache(
                                chat_output, generator_record, generator_key=f"run.outputs.{chat_output_name}"
                            ):
                                # there should be no extra spaces between adjacent chunks?
                                full_response += chunk
                                time.sleep(STREAMING_ANIMATION_TIME)
                                # Add a blinking cursor to simulate typing
                                message_placeholder.markdown(full_response + "▌")
                            message_placeholder.markdown(full_response)
                            response.output[chat_output_name] = full_response[prefix_length:]
                            post_process_dump_result(response, session_state_history, generator_record=generator_record)
                            return

            # generator in response has been fully consumed here
            resolved_outputs = post_process_dump_result(
                response, session_state_history, generator_record=generator_record
            )

            with container:
                render_message("assistant", resolved_outputs)

    image = Image.open(Path(__file__).parent / "logo.png")
    st.set_page_config(
        layout="wide",
        page_title=f"{flow_name}",
        page_icon=image,
        menu_items={
            "About": """
            HV Enigma Engine is a tool to give you the power to create any question instantly within mins.
            Fatser than Maggi :)

            Connect with me on: [LinkedIn](https://in.linkedin.com/in/harshavardhan-bajoria)
            """
        },
    )
    run_command()
    st.markdown(footer, unsafe_allow_html=True)
    # Set primary button color here since button color of the same form need to be identical in streamlit, but we only
    # need Run/Chat button to be blue.
    st.config.set_option("theme.primaryColor", "#0F6CBD")
    def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'

    color1 = "#0d3270"
    color2 = "#0fab7b"
    text = "HV Enigma Engine"
  
    left_co, cent_co, last_co = st.columns(3)
    with cent_co:
        st.image(str(Path(__file__).parent / "images/logo.png"), width=240)

    styled_text = gradient_text(text, color1, color2)
    st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    # st.title(flow_name)
    st.divider()
    st.chat_message("assistant").write("Hello, please input following flow inputs.")
    container = st.container()
    with container:
        show_conversation()

    with st.form(key="input_form", clear_on_submit=True):
        settings_path = os.path.join(os.path.dirname(__file__), "settings.json")
        if os.path.exists(settings_path):
            with open(settings_path, "r", encoding="utf-8") as file:
                json_data = json.load(file)
            environment_variables = list(json_data.keys())
            for environment_variable in environment_variables:
                secret_input="498b28b13a51434ab7c614bd4d961f20"
                if secret_input != "":
                    os.environ[environment_variable] = secret_input

        flow_inputs_params = {}
        for flow_input, (default_value, value_type) in flow_inputs.items():
            if value_type == "list":
                st.text(flow_input)
                from streamlit_quill import st_quill

                input = st_quill(
                    html=True,
                    toolbar=["image"],
                    key=flow_input,
                    placeholder="Please enter the list values and use the image icon to upload a picture. "
                    "Make sure to format each list item correctly with line breaks",
                )
            elif value_type == "image":
                input = st.file_uploader(label=flow_input)
            elif value_type == "string":
                input = st.text_input(label=flow_input, placeholder=default_value)
            else:
                input = st.text_input(label=flow_input, placeholder=default_value)
            flow_inputs_params.update({flow_input: copy(input)})

        cols = st.columns(7)
        submit_bt = cols[0].form_submit_button(label=label, type="primary")
        clear_bt = cols[1].form_submit_button(label="Clear")

        if submit_bt:
            with st.spinner("Loading..."):
                for flow_input, (default_value, value_type) in flow_inputs.items():
                    if value_type == "list":
                        input = parse_list_from_html(flow_inputs_params[flow_input])
                        flow_inputs_params.update({flow_input: copy(input)})
                    elif value_type == "image":
                        input = parse_image_content(
                            flow_inputs_params[flow_input],
                            flow_inputs_params[flow_input].type if flow_inputs_params[flow_input] else None,
                        )
                        flow_inputs_params.update({flow_input: copy(input)})
                submit(**flow_inputs_params)

        if clear_bt:
            with st.spinner("Cleaning..."):
                clear_chat()
                st.rerun()

    from docx import Document
    from io import BytesIO
    
    # Function to create a .docx file from text content
    def create_docx(content):
        doc = Document()
        doc.add_paragraph(content)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    # Read the .txt file, format the content, and save it as a .docx file
    with open("/workspaces/HV_Enigma/output/hvenigma.txt", "r") as file:
        file_content = file.read()
        file_content = file_content.replace("\\n", "\n").replace("\\t", "\t")
        file_content = file_content.replace("\\\"", "\"").replace("\\'", "'")
        # display the content in markdown format in the streamlit app as assisstant response
        file_content = format_markdown(file_content)
        st.markdown(file_content)

        # remove extra spaces from the markdown content
        file_content = file_content.replace("### ", "###")
        file_content = file_content.replace("#### ", "####")
        file_content = file_content.replace("## ", "##")
        file_content = file_content.replace("# ", "#")
        
        # create a .docx file from the content
        docx_buffer = create_docx(file_content)
    
    # Provide a download button for the .docx file
        st.download_button(
        label="Download Output",
        data=docx_buffer,
        file_name="hvenigma.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # # Add a download button for the output file
    # with open("/workspaces/HV_Enigma/output/hvenigma.txt", "rb") as file:
    #     # format the data to a visually appealing format as it is in markdown and contains \n and \t which should be replaced with actual new line and tabs.
    #     file = file.read().decode("utf-8")
    #     file = file.replace("\\n", "\n").replace("\\t", "\t")
    #     st.download_button(
    #         label="Download Output",
    #         data=file,
    #         file_name="hvenigma.txt",
    #         mime="text/plain"
    #     )


def resolve_flow_path(_from_config):
    if _from_config:
        result = Path(_from_config)
    else:
        result = Path(__file__).parent / "flow"
    if result.is_dir():
        os.chdir(result)
    else:
        os.chdir(result.parent)
    return result


if __name__ == "__main__":
    with open(Path(__file__).parent / "config.json", "r") as f:
        config = json.load(f)
        is_chat_flow = config["is_chat_flow"]
        chat_history_input_name = config["chat_history_input_name"]
        flow_path = resolve_flow_path(config["flow_path"])
        flow_name = config["flow_name"]
        flow_inputs = config["flow_inputs"]
        label = config["label"]
        is_streaming = config["is_streaming"]
        chat_output_name = config["chat_output_name"]

        footer = """
        <style>
        a:hover, a:active {
            color: red;
            background-color: transparent;
            text-decoration: underline;
        }
        .footer {
            position: fixed;
            left: 0;
            bottom: 0;
            width: 100%;
            background-color: white;
            color: black;
            text-align: center;
        }
        </style>
        <div class="footer">
            <p>Developed with ❤️ for <a style='display: inline; text-align: center;' href="https://www.unstop.com" target="_blank">Unstoppables</a></p>
        </div>
        """

    start()
    st.markdown(footer, unsafe_allow_html=True)